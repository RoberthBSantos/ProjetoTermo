from django.shortcuts import render, redirect, get_object_or_404
from .models import Produtos, ListaMaterial, Fornecedor, Grupos, DocFiles, Projeto
from .forms import FormularioContato, FormularioLista, FormularioFornecedor, NameForm, FormularioProjeto
import openpyxl
from openpyxl.styles import Font, colors, Alignment, Border, Side, PatternFill
import docx
from django.core.paginator import Paginator
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse, Http404
import os
from django.conf import settings
from django.core.files import File
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User


@login_required
def listar_contatos(request):
    busca = request.GET.get('pesquisa', None)

    if busca:
        # contatos = Contatos.objects.all()
        produtos_list = Produtos.objects.filter(nome__icontains=busca) or Produtos.objects.filter(
            fabricante__icontains=busca) or Produtos.objects.filter(modelo__icontains=busca)
    else:
        produtos_list = Produtos.objects.order_by('nome')

    paginator = Paginator(produtos_list, 10)

    page = request.GET.get('page')

    produtos = paginator.get_page(page)

    return render(request, 'contatos.html', {'contatos': produtos_list})


@login_required
def novo_projeto(request):
    form = FormularioProjeto(request.POST or None)
    if form.is_valid():
        projeto = form.save(commit= False)
        projeto.user= request.user
        projeto.save()

        return redirect('lista_projetos')

    return render(request, 'formulario_projeto.html', {'form': form})


@login_required
def novo_contato(request):
    form = FormularioContato(request.POST or None)

    if form.is_valid():
        form.save()
        return redirect('lista_contatos')

    return render(request, 'formulario_contato.html', {'form': form})


@login_required
def atualizar_projeto(request, id):
    projeto = get_object_or_404(Projeto, pk=id)
    form = FormularioProjeto(request.POST or None, instance=projeto)

    if form.is_valid():
        form.save()
        return redirect('lista_projetos')

    return render(request, 'formulario_projeto.html', {'form': form})


@login_required
def atualizar_contato(request, id):
    contato = get_object_or_404(Produtos, pk=id)
    form = FormularioContato(request.POST or None, instance=contato)

    if form.is_valid():
        form.save()
        return redirect('lista_contatos')

    return render(request, 'formulario_contato.html', {'form': form})


@login_required
def atualizar_fornecedor(request, id):
    fornecedor = get_object_or_404(Fornecedor, pk=id)
    form = FormularioFornecedor(request.POST or None, instance=fornecedor)

    if form.is_valid():
        form.save()
        return redirect('lista_fornecedor')

    return render(request, 'formulario_fornecedor.html', {'form': form})


def atualizar_prod_lista(request, id):
    produto = get_object_or_404(ListaMaterial, pk=id)
    form = FormularioLista(request.POST or None, instance=produto)
    projeto = produto.projeto
    if form.is_valid():
        form.save()
        return redirect('lista/id/', produto.projeto.id)

    return render(request, 'formulario_lista.html', {'form': form, 'projeto': projeto}, )


@login_required
def excluir_produto(request, id):
    produto = get_object_or_404(Produtos, id=id)
    form = FormularioContato(request.POST or None, request.FILES or None, instance=produto)
    if request.method == 'POST':
        produto.delete()
        return redirect('lista_contatos')

    # if request.method is not 'POST':
    #     post_delete=Contatos.objects.filter(id=id)
    #     post_delete.delete()
    #     return redirect('lista_contatos')

    return render(request, 'confirmar_delete_produto.html', {'produto': produto})


@login_required
def excluir_lista_produto(request, id):
    produto = get_object_or_404(ListaMaterial, id=id)
    form = FormularioContato(request.POST or None, request.FILES or None, instance=produto)
    id_lista = produto.projeto.id
    if request.method == 'POST':
        produto.delete()
        return redirect('lista/id/', id_lista)
    grupos = Grupos.objects.all()
    for grupo in grupos:
        print(grupo.nome)

    # if request.method is not 'POST':
    #     post_delete=Contatos.objects.filter(id=id)
    #     post_delete.delete()
    #     return redirect('lista_contatos')

    return render(request, 'confirmar_delete_produto.html', {'produto': produto})


@login_required
def excluir_prod_lista(request, id):
    produto = get_object_or_404(ListaMaterial, id=id)
    form = FormularioLista(request.POST or None)
    produtos = ListaMaterial.objects.all()
    contatos = Produtos.objects.all()

    if request.method is not 'POST':
        post_delete = ListaMaterial.objects.filter(id=id)
        post_delete.delete()
        produtos = ListaMaterial.objects.all()
        contatos = Produtos.objects.all()
        return render(request, 'formulario_lista.html', {'form': form, 'produtos': produtos, 'contatos': contatos})
    return render(request, 'formulario_lista.html', {'form': form, 'produtos': produtos, 'contatos': contatos})


@login_required
def novo_fornecedor(request):
    form = FormularioFornecedor(request.POST or None)
    convidar_usuario()

    if form.is_valid():
        form.save()
        return redirect('listar_fornecedor')
    return render(request, 'formulario_fornecedor.html', {'form': form})


@login_required
def listar_fornecedor(request):
    fornecedores = Fornecedor.objects.order_by('razao_social')

    return render(request, 'lista_fornecedor.html', {'fornecedores': fornecedores})


@login_required
def excluir_fornecedor(request, id):
    if request.method is not 'POST':
        forn_delete = Fornecedor.objects.filter(id=id)
        forn_delete.delete()
        return redirect('lista_fornecedor')


@login_required
def nova_lista(request, id):
    form = FormularioLista(request.POST or None)
    projeto = Projeto.objects.get(id=id)
    listas = ListaMaterial.objects.filter(projeto__id=id)

    if form.is_valid():

        lista = form.save(commit= False)
        lista.projeto = projeto
        for item in listas:
            if item.produto == lista.produto:
                return redirect('lista/id/', id)
        lista.save()


        return redirect('lista/id/', id)

    infra = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='INFRAESTRUTURA').order_by(
        'produto__nome')
    serv_infra = ListaMaterial.objects.filter(projeto=id).filter(
        produto__grupo__nome='SERVIÇOS DE INFRAESTRUTURA').order_by('produto__nome')
    fibra = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='FIBRA ÓPTICA').order_by(
        'produto__nome')
    ferragens = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='FERRAGENS E ACESSÓRIOS').order_by(
        'produto__nome')
    cabeamento = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='CABEAMENTO METÁLICO')
    racks = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='RACKS, GABINETES E ACESSÓRIOS')
    rede_eletrica = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='REDE ELÉTRICA')
    servicos_rede = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='SERVIÇOS DE REDE')
    rede_de_dados_e_energia = ListaMaterial.objects.filter(projeto=id).filter(
        produto__grupo__nome='REDE DE DADOS E ENERGIA')
    seguranca = ListaMaterial.objects.filter(projeto=id).filter(produto__grupo__nome='SEGURANÇA')

    projeto = Projeto.objects.get(id=id)

    return render(request, 'formulario_lista.html', {'form': form, 'infra': infra, 'serv_infra': serv_infra,
                                                     'fibra': fibra, 'ferragens': ferragens, 'cabeamento': cabeamento,
                                                     'racks': racks, 'rede_eletrica': rede_eletrica,
                                                     'servicos_rede': servicos_rede,
                                                     'rede_de_dados_e_energia': rede_de_dados_e_energia,
                                                     'seguranca': seguranca, 'projeto': projeto})


@login_required
def excluir_prod_lista(request, id):
    if request.method is not 'POST':
        id_lista = ListaMaterial.objects.filter(id=id).projeto
        ListaMaterial.objects.filter(id=id).delete()

        grupos = Grupos.objects.all()
        for grupo in grupos:
            print(grupo.nome)

        return redirect('lista/id/', id_lista)




@login_required
def gerar_xlsx(request, id):
    nome_doc = Projeto.objects.get(id=id).nome_projeto
    produtos = Produtos.objects.order_by('nome')
    lista = ListaMaterial.objects.filter(projeto=id, produto__grupo__nome='INFRAESTRUTURA')
    grupos = Grupos.objects.all()

    ##################  LISTAS DE MATERIAIS SEPARADAS POR GRUPO #########################


    wb = openpyxl.Workbook()
    cont = 2
    planilha = wb.active
    planilha.title = 'PREÇOS'

    # FONTES PARA FORMATAR AS CÉLULAS DO ARQUIVO
    ft_cabecalho = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
    ft_item = Font(name='Calibri', size=11)
    ft_item_negrito = Font(name='Arial', size=10, bold=True)
    ft_item_italico = Font(name='Arial', size=10, italic=True, color='505050')

    # ALINHAMENTOS
    alinhamento = Alignment(horizontal='center', vertical='center')
    alinhamentoEsquerda = Alignment(horizontal='left', vertical='center')
    alinhamentoDireita = Alignment(horizontal='right', vertical='center')

    # BORDAS
    fina = Side(border_style='thin', color='000000')
    bordaInferior = Border(bottom=fina)
    borda = Border(right=fina, bottom=fina)

    # PREENCHIMENTO DE CORES DAS CELULAS
    preenchimentoAzul = PatternFill('solid', fgColor='6495ED')
    preenchimentoVerde = PatternFill('solid', fgColor='00FF7F')
    preenchimentoCinza = PatternFill('solid', fgColor='DDDDDD')
    preenchimentoGrupo = PatternFill('solid', fgColor='A9A9A9')
    preenchimentoAzulClaro = PatternFill('solid', fgColor='dbe5f1')

    # COMEÇANDO A ESCREVER O ARQUIVO XLSX
    linha = 1
    i_grupo = 1
    i = 1

    for grupo in grupos:
        if grupo.nome == 'INFRAESTRUTURA':
            planilha.merge_cells('A' + str(linha) + ':N' + str(linha))
            planilha['A' + str(linha)] = 'UPI (UNIDADE DE PLANTA DE INFRAESTRUTURA)'
            planilha['A' + str(linha)].alignment = alinhamento
            planilha['A' + str(linha)].fill = preenchimentoAzul
            planilha['A' + str(linha)].font = ft_cabecalho

            linha += 1
        if grupo.nome == 'FIBRA ÓPTICA':
            planilha.merge_cells('A' + str(linha) + ':N' + str(linha))
            planilha['A' + str(linha)] = 'UPR (UNIDADE DE PLANTA DE REDE)'
            planilha['A' + str(linha)].alignment = alinhamento
            planilha['A' + str(linha)].fill = preenchimentoAzul
            planilha['A' + str(linha)].font = ft_cabecalho

            linha += 1
        if grupo.nome == 'REDE DE DADOS E ENERGIA':
            planilha.merge_cells('A' + str(linha) + ':N' + str(linha))
            planilha['A' + str(linha)] = 'UPE (UNIDADE DE PLANTA DE EQUIPAMENTO)'
            planilha['A' + str(linha)].alignment = alinhamento
            planilha['A' + str(linha)].fill = preenchimentoAzul
            planilha['A' + str(linha)].font = ft_cabecalho

            linha += 1
        lista_grupo = ListaMaterial.objects.filter(projeto=id, produto__grupo__nome=grupo.nome)
        if len(lista_grupo) > 0:
            planilha.merge_cells('A' + str(linha) + ':N' + str(linha))
            planilha['A' + str(linha)] = 'GRUPO ' + str(i_grupo) +' '+ grupo.nome
            planilha['A' + str(linha)].alignment = alinhamento
            planilha['A' + str(linha)].fill = preenchimentoGrupo
            planilha['A' + str(linha)].font = ft_cabecalho

            linha += 1

            planilha['A' + str(linha)].fill = preenchimentoGrupo
            planilha['B' + str(linha)].fill = preenchimentoGrupo
            planilha['C' + str(linha)].fill = preenchimentoGrupo
            planilha['D' + str(linha)].fill = preenchimentoGrupo
            planilha['E' + str(linha)].fill = preenchimentoGrupo
            planilha['F' + str(linha)].fill = preenchimentoGrupo
            planilha['G' + str(linha)].fill = preenchimentoGrupo
            planilha['H' + str(linha)].fill = preenchimentoGrupo
            planilha['I' + str(linha)].fill = preenchimentoGrupo
            planilha['J' + str(linha)].fill = preenchimentoGrupo
            planilha['K' + str(linha)].fill = preenchimentoGrupo
            planilha['L' + str(linha)].fill = preenchimentoGrupo
            planilha['M' + str(linha)].fill = preenchimentoGrupo
            planilha['N' + str(linha)].fill = preenchimentoGrupo

            planilha['A' + str(linha)].font = ft_cabecalho
            planilha['B' + str(linha)].font = ft_cabecalho
            planilha['C' + str(linha)].font = ft_cabecalho
            planilha['D' + str(linha)].font = ft_cabecalho
            planilha['E' + str(linha)].font = ft_cabecalho
            planilha['F' + str(linha)].font = ft_cabecalho
            planilha['G' + str(linha)].font = ft_cabecalho
            planilha['H' + str(linha)].font = ft_cabecalho
            planilha['I' + str(linha)].font = ft_cabecalho
            planilha['J' + str(linha)].font = ft_cabecalho
            planilha['K' + str(linha)].font = ft_cabecalho
            planilha['L' + str(linha)].font = ft_cabecalho
            planilha['M' + str(linha)].font = ft_cabecalho
            planilha['N' + str(linha)].font = ft_cabecalho

            planilha['A' + str(linha)].border = borda
            planilha['B' + str(linha)].border = borda
            planilha['C' + str(linha)].border = borda
            planilha['D' + str(linha)].border = borda
            planilha['E' + str(linha)].border = borda
            planilha['F' + str(linha)].border = borda
            planilha['G' + str(linha)].border = borda
            planilha['H' + str(linha)].border = borda
            planilha['I' + str(linha)].border = borda
            planilha['J' + str(linha)].border = borda
            planilha['K' + str(linha)].border = borda
            planilha['L' + str(linha)].border = borda
            planilha['M' + str(linha)].border = borda
            planilha['N' + str(linha)].border = borda

            planilha['A' + str(linha)].alignment = alinhamento
            planilha['B' + str(linha)].alignment = alinhamento
            planilha['C' + str(linha)].alignment = alinhamento
            planilha['D' + str(linha)].alignment = alinhamento
            planilha['E' + str(linha)].alignment = alinhamento
            planilha['F' + str(linha)].alignment = alinhamento
            planilha['G' + str(linha)].alignment = alinhamento
            planilha['H' + str(linha)].alignment = alinhamento
            planilha['I' + str(linha)].alignment = alinhamento
            planilha['J' + str(linha)].alignment = alinhamento
            planilha['K' + str(linha)].alignment = alinhamento
            planilha['L' + str(linha)].alignment = alinhamento
            planilha['M' + str(linha)].alignment = alinhamento
            planilha['N' + str(linha)].alignment = alinhamento

            planilha['A' + str(linha)] = 'FABRICANTE'
            planilha['B' + str(linha)] = 'MODELO'
            planilha['C' + str(linha)] = 'DESCRIÇÃO'
            planilha['D' + str(linha)] = 'UND'
            planilha['E' + str(linha)] = 'PONTOS'
            planilha['F' + str(linha)] = 'CUSTO PROD.'
            planilha['G' + str(linha)] = 'VENDA PROD.'
            planilha['H' + str(linha)] = '∆T INF.'
            planilha['I' + str(linha)] = '∆T SUP.'
            planilha['J' + str(linha)] = 'SERV. TERC.'
            planilha['K' + str(linha)] = 'SERVIÇO'
            planilha['L' + str(linha)] = 'QTD'
            planilha['M' + str(linha)] = 'PREÇO'
            planilha['N' + str(linha)] = 'SUBTOTAL'

            linha += 1

            indice = linha

            for item in lista_grupo:

                planilha['A' + str(linha)] = item.produto.fabricante
                planilha['B' + str(linha)] = item.produto.modelo
                planilha['C' + str(linha)] = item.produto.nome
                planilha['D' + str(linha)] = item.produto.unidade
                planilha['E' + str(linha)] = item.pontos
                planilha['F' + str(linha)] = item.produto.valor_de_compra
                planilha['F' + str(linha)].number_format = '_-R$ * ##,##0.00_-'
                planilha['G' + str(linha)].number_format = '_-R$ * ##,##0.00_-'
                # planilha['G' + str(cont)] = format(produto.data, "%d/%m/%Y")
                planilha['G' + str(linha)] = item.custo_produto
                planilha['H' + str(linha)] = item.produto.tempo_de_instalacao
                if item.produto.tempo_de_sup > 0:
                    planilha['I' + str(linha)] = item.produto.tempo_de_sup
                if item.produto.valor_de_terceiros > 0:
                    planilha['J' + str(linha)] = item.produto.valor_de_terceiros
                    planilha['J' + str(linha)].number_format = '_-R$ * ##,##0.00_-'
                planilha['K' + str(linha)] = item.custo_servico
                planilha['K' + str(linha)].number_format = '_-R$ * ##,##0.00_-'
                planilha['L' + str(linha)] = item.quantidade
                planilha['M' + str(linha)] = item.custo_venda
                planilha['M' + str(linha)].number_format = '_-R$ * ##,##0.00_-'
                planilha['N' + str(linha)] = '= M' + str(linha) + '*L' + str(linha)
                planilha['N' + str(linha)].number_format = '_-R$ * ##,##0.00_-'

                planilha['A' + str(linha)].font = ft_item
                planilha['B' + str(linha)].font = ft_item
                planilha['C' + str(linha)].font = ft_item
                planilha['D' + str(linha)].font = ft_item
                planilha['E' + str(linha)].font = ft_item
                planilha['F' + str(linha)].font = ft_item
                planilha['G' + str(linha)].font = ft_item
                planilha['I' + str(linha)].font = ft_item
                planilha['J' + str(linha)].font = ft_item
                planilha['K' + str(linha)].font = ft_item
                planilha['L' + str(linha)].font = ft_item
                planilha['M' + str(linha)].font = ft_item
                planilha['N' + str(linha)].font = ft_item

                planilha['A' + str(linha)].border = borda
                planilha['B' + str(linha)].border = borda
                planilha['C' + str(linha)].border = borda
                planilha['D' + str(linha)].border = borda
                planilha['E' + str(linha)].border = borda
                planilha['F' + str(linha)].border = borda
                planilha['G' + str(linha)].border = borda
                planilha['H' + str(linha)].border = borda
                planilha['I' + str(linha)].border = borda
                planilha['J' + str(linha)].border = borda
                planilha['K' + str(linha)].border = borda
                planilha['L' + str(linha)].border = borda
                planilha['M' + str(linha)].border = borda
                planilha['N' + str(linha)].border = borda

                planilha['A' + str(linha)].alignment = alinhamento
                planilha['B' + str(linha)].alignment = alinhamento
                planilha['C' + str(linha)].alignment = alinhamentoEsquerda
                planilha['D' + str(linha)].alignment = alinhamento
                planilha['E' + str(linha)].alignment = alinhamento
                planilha['F' + str(linha)].alignment = alinhamento
                planilha['G' + str(linha)].alignment = alinhamento
                planilha['H' + str(linha)].alignment = alinhamento
                planilha['I' + str(linha)].alignment = alinhamento
                planilha['K' + str(linha)].alignment = alinhamento
                planilha['L' + str(linha)].alignment = alinhamento
                planilha['M' + str(linha)].alignment = alinhamento
                planilha['N' + str(linha)].alignment = alinhamento

                if linha % 2 == 0:
                    planilha['A' + str(linha)].fill = preenchimentoCinza
                    planilha['B' + str(linha)].fill = preenchimentoCinza
                    planilha['C' + str(linha)].fill = preenchimentoCinza
                    planilha['D' + str(linha)].fill = preenchimentoCinza
                    planilha['E' + str(linha)].fill = preenchimentoCinza
                    planilha['F' + str(linha)].fill = preenchimentoCinza
                    planilha['G' + str(linha)].fill = preenchimentoCinza
                    planilha['H' + str(linha)].fill = preenchimentoCinza
                    planilha['I' + str(linha)].fill = preenchimentoCinza
                    planilha['J' + str(linha)].fill = preenchimentoCinza
                    planilha['K' + str(linha)].fill = preenchimentoCinza
                    planilha['L' + str(linha)].fill = preenchimentoCinza
                    planilha['M' + str(linha)].fill = preenchimentoCinza
                    planilha['N' + str(linha)].fill = preenchimentoCinza

                else:
                    planilha['A' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['B' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['C' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['D' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['E' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['F' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['G' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['H' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['I' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['J' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['K' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['L' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['M' + str(linha)].fill = preenchimentoAzulClaro
                    planilha['N' + str(linha)].fill = preenchimentoAzulClaro

                linha += 1
            planilha['N' + str(linha)] = '= SUM(N' + str(indice) + ':N' + str(linha - 1) + ')'
            planilha['N' + str(linha)].number_format = '_-R$ * ##,##0.00_-'
            planilha['N' + str(linha)].font = ft_item_negrito
            planilha['N' + str(linha)].alignment = alinhamento
            planilha['N' + str(linha)].fill = preenchimentoVerde

            linha += 1
            i_grupo += 1


    planilha.column_dimensions["A"].width = 20.0
    planilha.column_dimensions["B"].width = 20.0
    planilha.column_dimensions["C"].width = 101.0
    planilha.column_dimensions["D"].width = 10.0
    planilha.column_dimensions["E"].width = 12.0
    planilha.column_dimensions["F"].width = 15.0
    planilha.column_dimensions["G"].width = 15.0
    planilha.column_dimensions["H"].width = 8.0
    planilha.column_dimensions["K"].width = 15.0
    planilha.column_dimensions["J"].width = 12.0
    planilha.column_dimensions["L"].width = 15.0
    planilha.column_dimensions["M"].width = 15.0
    planilha.column_dimensions["N"].width = 17.0



    wb.save('documents/documents/media/Planilha ' + nome_doc + '.xlsx')
    gerar_planilha(nome_doc)
    return redirect('listar_downloads')


@login_required
def gerar_docx(request, id):
    nome_doc = Projeto.objects.get(id=id).nome_projeto
    produtos = Produtos.objects.order_by('nome')
    lista = ListaMaterial.objects.filter(projeto=id).order_by('produto__nome')
    grupos = Grupos.objects.all()
    cgrupo = 1

    doc = docx.Document()

    doc.add_paragraph('ANEXO A – ESPECIFICAÇÕES TÉCNICAS').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[0].runs[0].font.size = Pt(16)
    doc.paragraphs[0].runs[0].font.bold = True
    doc.paragraphs[0].runs[0].font.name = 'Calibri'
    doc.add_paragraph(
        'TODOS OS PRODUTOS OFERTADOS DEVERÃO TER, EM SUA COMPOSIÇÃO DE CUSTOS, OS VALORES REFERENTE A INSTALAÇÃO.').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.paragraphs[0].runs[0].font.size = Pt(16)
    doc.paragraphs[1].runs[0].font.size = Pt(11)
    doc.paragraphs[1].runs[0].font.name = 'Calibri'
    doc.add_paragraph('UPI – UNIDADE DE PONTO DE INFRAESTRUTURA').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.paragraphs[2].runs[0].font.size = Pt(13)
    doc.paragraphs[2].runs[0].font.bold = True
    doc.paragraphs[2].runs[0].font.name = 'Calibri'
    doc.paragraphs[2].runs[0].font.color.rgb = RGBColor(79, 129, 189)

    for grupo in grupos:
        if group_check(grupo.nome):
            contador = 1
            index_group = doc.add_paragraph()
            titulo = ('GRUPO ', str(cgrupo), ' ', grupo.nome)
            nome = index_group.add_run(titulo)
            nome.font.bold = True
            nome.font.color.rgb = RGBColor(79, 129, 189)
            for item in lista:
                for produto in produtos:
                    if item.produto.id == produto.id and str(produto.grupo) == grupo.nome:

                        run = doc.add_paragraph().add_run()

                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Calibri'
                        font.size = docx.shared.Pt(11)
                        descricao = produto.descricao
                        descricao = descricao.splitlines()
                        titulo = (str(cgrupo) + '.' + str(contador) + ' ' + produto.nome)

                        paragrafo = doc.add_paragraph()
                        nome = paragrafo.add_run((titulo))
                        nome.font.bold = True
                        nome.font.color.rgb = RGBColor(79, 129, 189)

                        contador += 1
                        for topico in descricao:
                            doc.add_paragraph(
                                (topico), style='List Bullet'
                            )
            cgrupo += 1

    doc.save('documents/documents/media/Anexos ' + nome_doc + '.docx')
    gerar_doc(nome_doc)

    return redirect('listar_downloads')


# FUNÇÃO PARA CHECAR SE HÁ ITENS NO GRUPO SOLICITADO.
def group_check(grupo):
    produtos = Produtos.objects.order_by('nome')
    lista = ListaMaterial.objects.order_by('produto__nome')
    for item in lista:
        for produto in produtos:
            if item.produto.id == produto.id and str(produto.grupo) == grupo:
                return True

    return False


def gerar_doc(nome):
    f = File(open(os.path.join(settings.MEDIA_ROOT, 'documents/media/Anexos ' + nome + '.docx'), 'rb'))
    doc = DocFiles()
    doc.docupload = f

    doc.title = 'Anexos ' + nome

    doc.save(nome)


def gerar_planilha(nome):
    f = File(open(os.path.join(settings.MEDIA_ROOT, 'documents/media/Planilha ' + nome + '.xlsx'), 'rb'))
    doc = DocFiles()
    doc.docupload = f

    doc.title = 'Planilha ' + nome

    doc.save(nome)


def download_doc(path):
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="aplication/docupload")
            response['Content-Disposition'] = 'inline;  filename=' + os.path.basename(file_path)
            return response
        raise Http404


def listar_download(request):
    files = {'files': DocFiles.objects.all()}
    return render(request, 'download_list.html', files)


def get_name_docx(request, id):
    if request.method == 'POST':
        form = NameForm(request.POST)
        if form.is_valid():
            nome = form.cleaned_data['project_name']
            gerar_docx(nome)
            return redirect('listar_downloads')

    else:
        form = NameForm()
    return render(request, 'formulario_docs.html', {'form': form})


def get_name_xlsx(request):
    if request.method == 'POST':
        form = NameForm(request.POST)

        if form.is_valid():
            nome = form.cleaned_data['project_name']
            gerar_xlsx(nome)
            return redirect('listar_downloads')

    else:
        form = NameForm()

    return render(request, 'formulario_docs.html', {'form': form})


@login_required
def delete_doc(request, id):
    documento = get_object_or_404(DocFiles, id=id)
    try:
        os.remove('documents/documents/media/' + documento.title + '.docx')
    except:
        os.remove('documents/documents/media/' + documento.title + '.xlsx')
    os.remove(str(documento.docupload))
    DocFiles.objects.filter(id=id).delete()
    return redirect('listar_downloads')


def limpar_lista(request):
    lista = ListaMaterial.objects.all()

    for item in lista:
        item.delete()

    return redirect('adicionar_lista')


@login_required
def listar_projetos(request):
    busca = request.GET.get('pesquisa', None)
    teste = Projeto.objects.all()

    if busca:
        # contatos = Contatos.objects.all()
        projetos = {'projetos': Projeto.objects.filter(nome_projeto__icontains=busca,user= request.user)}
    else:
        projetos = {'projetos': Projeto.objects.filter(user=request.user) | Projeto.objects.filter(convidados__id=request.user.id)}

    return render(request, 'index.html', projetos)


def vincular_projeto(id):

    projeto = Projeto.objects.get(id=id)
    lista = ListaMaterial.objects.get(projeto=None)
    lista.projeto = projeto
    lista.save()

def deletar_projeto(request, id):
    projeto = get_object_or_404(Projeto, id=id)

    if request.method == 'POST':
        projeto.delete()

        return redirect('lista_projetos')

    return render(request, 'confirmar_delete_produto.html', {'projeto': projeto})

def convidar_usuario():
    usuarios = User.objects.all()
    for usuario  in usuarios:
        print(usuario.first_name +' '+ usuario.last_name)

@login_required
def get_perfil_logado(request):
    return request.user.perfil
